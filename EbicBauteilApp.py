# EbicBauteilApp.py
from __future__ import annotations

import os
import io
import re
import json
import base64
import zipfile
import sqlite3
from datetime import datetime
from dataclasses import dataclass
from typing import List, Optional, Dict, Any

import pandas as pd
from PIL import Image, ExifTags
import streamlit as st

from pathlib import Path

# ---- Session-State Defaults ----
def init_state():
    defaults = {
        "gallery_selected": [],      # oder None, je nach Logik
        "gallery_items": [],
        "page": "Projekt erstellen", # Standard-Seite
        "prompt_text": "",
        "ki_feedback": "",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()  # <-- direkt nach den Imports aufrufen


# ========= BASISPFAD / LOGOS / DB / CSV =========
BASE_DIR = os.path.join(os.path.dirname(__file__), "Bilder")

LOGO_PATH_SIDEBAR = os.path.join(BASE_DIR, "EmchBergerLogo.png")
LOGO_PATH_PROJEKTE = os.path.join(BASE_DIR, "apartment.png")
LOGO_PATH_PROJEKTERSTELLEN = os.path.join(BASE_DIR, "engineering.png")
LOGO_PATH_EXPORT = os.path.join(BASE_DIR, "move.png")
LOGO_PATH_EINSTELLUNG = os.path.join(BASE_DIR, "settings.png")
EBKPH_CSV_PATH = os.path.join(BASE_DIR, "eBKPH_Gesamttabelle.csv")
DB_PATH = os.path.join(os.path.dirname(__file__), "projects.db")

# ========= DB: Tabelle sicherstellen =========
def ensure_db():
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("""
    CREATE TABLE IF NOT EXISTS projects (
        name TEXT PRIMARY KEY,
        nr TEXT,
        addr TEXT,
        plz TEXT,
        ort TEXT,
        meta_json TEXT,
        data_json TEXT,
        images_zip BLOB,
        thumbs_zip BLOB,
        created_at TEXT,
        report_path TEXT
    );
    """)
    # NEU: eBKP-H Auswahl-Batches
    cur.execute("""
    CREATE TABLE IF NOT EXISTS ebkh_selection (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        code TEXT NOT NULL,
        batch_ts TEXT NOT NULL
    );
    """)
    con.commit()
    con.close()

ensure_db()

# ========= App-Konstanten =========
APP_TITLE = "KI-Bauteilbeschrieb"
THUMB_MAX = (256, 256)
DEFAULT_GESCHOSSE = ["UG", "EG", "OG", "DG", "Unbekannt"]
DEFAULT_KLASSEN = [
    "Elektro / Verteilung / Kabel",
    "Sanitär / Rohrleitungen",
    "HLK / Klima / Lüftung",
    "Fenster / Türen / Außenbauteile",
    "Rohbau / Mauerwerk / Decke",
    "Sicherheit / Brandschutz",
    "Sonstiges / Umfeld",
]
TABLE_COLUMNS_ORDER = [
    "Miniatur", "Bild original", "Bild umbenennen", "Zeit", "Geschoss", "Kategorie",
    "Bauteile", "Besonderheiten", "Risiken", "Empfehlung", "Pfad"
]

# ========= EXIF-Zeit lesen =========
EXIF_DATETIME_TAGS = {
    ExifTags.TAGS.get(k) for k in ExifTags.TAGS
    if str(ExifTags.TAGS.get(k, "")).startswith("DateTime")
}

def read_exif_datetime(img: Image.Image) -> Optional[datetime]:
    try:
        ex = img.getexif()
        if not ex:
            return None
        for k, v in ex.items():
            tag = ExifTags.TAGS.get(k, str(k))
            if tag in EXIF_DATETIME_TAGS and isinstance(v, str) and ":" in v:
                vv = v.replace(":", "-", 2)
                try:
                    return datetime.fromisoformat(vv)
                except Exception:
                    pass
    except Exception:
        return None
    return None

# ========= Datensatz-Modell =========
@dataclass
class FotoRecord:
    preview: bytes
    original_name: str
    rename_to: str
    zeit: str
    geschoss: str
    kategorie: str
    bauteile: str
    besonderheiten: str
    risiken: str
    empfehlung: str
    pfad: str

    def to_row(self) -> Dict[str, str]:
        img_b64 = base64.b64encode(self.preview).decode("ascii")
        return {
            "Miniatur": f"data:image/jpeg;base64,{img_b64}",
            "Bild original": self.original_name,
            "Bild umbenennen": self.rename_to,
            "Zeit": self.zeit,
            "Geschoss": self.geschoss,
            "Kategorie": self.kategorie,
            "Bauteile": self.bauteile,
            "Besonderheiten": self.besonderheiten,
            "Risiken": self.risiken,
            "Empfehlung": self.empfehlung,
            "Pfad": self.pfad,
        }

# ========= Streamlit Setup =========
st.set_page_config(page_title=APP_TITLE, layout="wide")


# Header mit Icon + Text
col_icon, col_text = st.columns([0.1, 0.9])
with col_icon:
    logo_path = os.path.join(BASE_DIR, "AI-Logo.jpg")  # dein Logo
    if os.path.exists(logo_path):
        st.image(logo_path, width=300,)  # oder width="stretch"

with col_text:
    st.markdown(
        "<div style='display:flex;align-items:center;height:48px;"
        "font-size:50px;font-weight:700;'>Zustandsanalyse mit KI</div>",
        unsafe_allow_html=True
    )

# ========= Globale Styles =========
def inject_css(hover_color: str, hover_text: str):
    st.markdown(
        f"""
        <style>
          :root {{
            --hover-color: {hover_color};
            --hover-text: {hover_text};
          }}

          /* ========== Buttons & Tabs (Hover: BG = hover-color, Text = weiß) ========== */
          .stButton > button:hover,
          .stDownloadButton > button:hover,
          button[kind]:hover,
          [role="button"]:hover,
          .stTabs button[role="tab"]:hover,
          .nav-row button:hover {{
            background-color: var(--hover-color) !important;
            border-color: var(--hover-color) !important;
            color: var(--hover-text) !important;
          }}

          /* auch bei :focus-visible, falls per Tastatur */
          .stButton > button:focus-visible,
          .stDownloadButton > button:focus-visible,
          button[kind]:focus-visible,
          .stTabs button[role="tab"]:focus-visible,
          .nav-row button:focus-visible {{
            outline: 2px solid var(--hover-color) !important;
            outline-offset: 2px !important;
          }}

          /* ========== Links ========== */
          a:hover {{
            color: var(--hover-color) !important;
          }}

          /* ========== Inputs (Hover: Rahmen/Glow statt Vollfläche) ========== */
          input[type="text"]:hover,
          input[type="number"]:hover,
          input[type="email"]:hover,
          input[type="password"]:hover,
          input[type="search"]:hover,
          textarea:hover,
          select:hover {{
            border-color: var(--hover-color) !important;
            box-shadow: 0 0 0 1px var(--hover-color) inset !important;
          }}

          /* Optional: Label-Farbe bei Hover auf Radio/Checkbox/Select */
          label:hover {{
            color: var(--hover-color) !important;
          }}
          [role="radio"]:hover label,
          [role="checkbox"]:hover label {{
            color: var(--hover-color) !important;
          }}

          /* ========== Tabellen / DataFrames / Data Editor ========== */
          /* einfache Tabellen-Zeilen */
          table tbody tr:hover td,
          table tbody tr:hover th {{
            background-color: rgba(0, 176, 202, 0.12) !important;
          }}
          /* Streamlit DataFrame/DataEditor (breit gefasst) */
          .stDataFrame tbody tr:hover td,
          .stDataFrame tbody tr:hover th {{
            background-color: rgba(0, 176, 202, 0.12) !important;
          }}
        </style>
        """,
        unsafe_allow_html=True
    )


# ========= Toast-Fallback =========
if not hasattr(st, "toast"):
    def _toast(msg: str, icon: Optional[str] = None):
        if icon == "✅": st.success(msg)
        elif icon == "❌": st.error(msg)
        elif icon == "⚠️": st.warning(msg)
        else: st.info(msg)
    st.toast = _toast

# ========= Persistente Konfiguration =========
CONFIG_PATH = os.path.join(os.path.dirname(__file__), "config.json")

def load_config() -> dict:
    base = {
        "geschosse": DEFAULT_GESCHOSSE,
        "klassen": DEFAULT_KLASSEN,
        "ki_provider": "OpenAI (ChatGPT)",
        "ki_model": "gpt-4o",
        "hover_color": "#00b0ca",
        "hover_text_color": "#ffffff",
    }
    if os.path.exists(CONFIG_PATH):
        try:
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
            # Nur harmlose Sachen übernehmen
            base.update({
                "geschosse": data.get("geschosse", base["geschosse"]),
                "klassen": data.get("klassen", base["klassen"]),
                "ki_provider": data.get("ki_provider", base["ki_provider"]),
                "ki_model": data.get("ki_model", base["ki_model"]),
                "hover_color": data.get("hover_color", base["hover_color"]),
                "hover_text_color": data.get("hover_text_color", base["hover_text_color"]),
            })
        except Exception:
            pass
    return base

def save_config(cfg: dict):
    try:
        # Sicherheitskopie ohne API-Keys
        safe_cfg = {k: v for k, v in cfg.items() if k != "api_keys"}
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(safe_cfg, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


# ========= Session-Init =========
if "page" not in st.session_state: st.session_state.page = "Projektspeicher"
if "last_preview_pdf" not in st.session_state: st.session_state.last_preview_pdf = None

if "config" not in st.session_state: st.session_state.config = load_config()
inject_css(
    st.session_state.config.get("hover_color", "#00b0ca"),
    st.session_state.config.get("hover_text_color", "#ffffff"),
)
if "filters" not in st.session_state: st.session_state.filters = {"geschoss": [], "kategorie": [], "text": ""}
if "upload_token" not in st.session_state: st.session_state.upload_token = 0
if "undo_stack" not in st.session_state: st.session_state.undo_stack = []

# ========= Helper: Icons im Sidebar =========
def _icon_data_uri(path: str) -> Optional[str]:
    if not (path and os.path.exists(path)): return None
    ext = os.path.splitext(path)[1].lower()
    mime = "image/png" if ext == ".png" else "image/jpeg"
    try:
        with open(path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("ascii")
        return f"data:{mime};base64,{b64}"
    except Exception:
        return None

def nav_item(label: str, icon_path: str, target_page: str):
    active = (st.session_state.page == target_page)
    st.markdown('<div class="nav-row">', unsafe_allow_html=True)
    try:
        c1, c2 = st.columns([0.18, 0.82], vertical_alignment="top")
    except TypeError:
        c1, c2 = st.columns([0.18, 0.82])
    with c1:
        src = _icon_data_uri(icon_path)
        if src:
            st.markdown(f'<div class="nav-icon-box"><img class="nav-icon-img" src="{src}" alt=""></div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="nav-icon-box">❔</div>', unsafe_allow_html=True)
    with c2:
        if st.button(label, key=f"nav_{target_page}", width="stretch",
                     type=("primary" if active else "secondary")):
            st.session_state.page = target_page
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# ======== KI: Custom-Prompt Helpers ========
def _get_active_ki_settings():
    # 1) Anbieter/Modell wie bisher aus der Config lesen
    cfg = st.session_state.config
    provider = cfg.get("ki_provider", "OpenAI (ChatGPT)")
    model = cfg.get("ki_model", "gpt-4o")

    # 2) API-Key bevorzugt aus Secrets/Umgebung (sicher!)
    secret_key = (st.secrets.get("OPENAI_API_KEY")
                  or os.environ.get("OPENAI_API_KEY")
                  or "")

    if secret_key:
        api_key = secret_key
    else:
        # 3) Fallback: alter Weg (nur falls noch irgendwo in config.json)
        api_keys = (cfg.get("api_keys") or {}).get(provider, {})
        api_key = api_keys.get(model) or api_keys.get("default") or ""

    return provider, model, api_key



def _apply_custom_result_to_df(df: pd.DataFrame, results: dict):
    """Schreibt die Felder aus 'results' in die DataFrame-Zeilen."""
    if df is None or df.empty:
        return df
    base = df.copy()
    for name, payload in results.items():
        mask = (base["Bild original"] == name) if "Bild original" in base.columns else None
        if mask is not None and mask.any():
            for col in ["Bauteile", "Besonderheiten", "Risiken", "Empfehlung"]:
                if col in base.columns and payload.get(col) is not None:
                    base.loc[mask, col] = str(payload.get(col, "")).strip()
    return base


def run_custom_prompt_openai(prompt: str, model: str, api_key: str, targets: list, df: pd.DataFrame) -> dict:
    """
    Ruft OpenAI (Chat Completions) auf und bittet um JSON mit
    Bauteile/Besonderheiten/Risiken/Empfehlung pro Bild.
    """
    if not targets:
        return {}

    # Kontext aus Tabelle
    def _row_dict(name):
        try:
            row = df[df["Bild original"] == name].head(1).to_dict(orient="records")
            return row[0] if row else {"Bild original": name}
        except Exception:
            return {"Bild original": name}

    system_msg = (
        "Du bist ein Assistent für Zustandsanalysen in Gebäuden. "
        "Antworte ausschließlich mit JSON der Form:\n"
        "{\n"
        '  "<bildname>": {\n'
        '    "Bauteile": "...",\n'
        '    "Besonderheiten": "...",\n'
        '    "Risiken": "...",\n'
        '    "Empfehlung": "..."\n'
        "  }\n"
        "}"
    )

    user_payload = {
        "hinweis": "Überarbeite/ergänze die folgenden Datensätze.",
        "prompt": prompt,
        "ziele": [{"Bild original": t, "context": _row_dict(t)} for t in targets],
    }

    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)},
            ],
            temperature=0.2,
        )
        content = resp.choices[0].message.content.strip()
    except Exception as e:
        st.toast(f"OpenAI-Aufruf fehlgeschlagen: {e}", icon="❌")
        return {}

    # JSON parsen
    try:
        if content.startswith("```"):
            content = content.strip("`")
            if content.lower().startswith("json"):
                content = content[4:].lstrip()
        data = json.loads(content)
        if not isinstance(data, dict):
            raise ValueError("Antwort ist kein JSON-Objekt")
        results = {}
        for name in targets:
            block = data.get(name) or {}
            results[name] = {
                "Bauteile": block.get("Bauteile", ""),
                "Besonderheiten": block.get("Besonderheiten", ""),
                "Risiken": block.get("Risiken", ""),
                "Empfehlung": block.get("Empfehlung", ""),
            }
        return results
    except Exception as e:
        st.toast(f"Antwort nicht als JSON lesbar: {e}", icon="⚠️")
        return {}

# ========= Sidebar =========

with st.sidebar:
    if os.path.exists(LOGO_PATH_SIDEBAR):
        c1, c2 = st.columns([0.25, 0.75])
        with c1: st.image(LOGO_PATH_SIDEBAR, width=48)
        with c2: st.markdown("<div style='display:flex;align-items:center;height:48px;font-size:20px;font-weight:700;'>KI-Bauteilbeschrieb</div>", unsafe_allow_html=True)
    st.divider()
    st.markdown("### Navigation")
    nav_item("Projekte", LOGO_PATH_PROJEKTE, "Projektspeicher")
    nav_item("Projekt erstellen", LOGO_PATH_PROJEKTERSTELLEN, "Projekt")
    nav_item("Export", LOGO_PATH_EXPORT, "Export")
    nav_item("Einstellung", LOGO_PATH_EINSTELLUNG, "Einstellung")
    st.divider()

    if st.session_state.page == "Projekt":
        st.markdown("### Filtern")
        cfg = st.session_state.config
        st.session_state.filters["geschoss"] = st.multiselect("Geschosse", options=cfg["geschosse"])
        st.session_state.filters["kategorie"] = st.multiselect("Kategorien", options=cfg["klassen"])
        st.session_state.filters["text"] = st.text_input("Textfilter", value="", placeholder="z. B. Dose, Trennschalter…")
        st.divider()


    # ======== Custom-Prompt für KI (nur auf Seite "Projekt") ========
    if st.session_state.page == "Projekt":
        st.markdown("### KI – benutzerdefinierte Prompt")

        prov, mdl, api_key = _get_active_ki_settings()
        st.caption(
            f"Verwendeter Anbieter/Modell: **{prov} / {mdl}** "
            f"{'(Key vorhanden)' if api_key else '(kein API-Key hinterlegt!)'}"
        )

        # Zielauswahl: Bildnamen aus aktueller Tabelle
        base_df = st.session_state.get("base_df", None)
        img_options = (base_df["Bild original"].dropna().tolist()
                    if (base_df is not None and not base_df.empty and "Bild original" in base_df.columns)
                    else [])
        targets = st.multiselect("Auf diese Bilder anwenden", options=img_options)

        custom_prompt = st.text_area(
            "Dein Prompt an die KI",
            placeholder="Beschreibe hier, wie die Analyse geändert/ergänzt werden soll …",
            height=120
        )

        run_disabled = (not custom_prompt.strip()) or (not targets)
        if st.button("Prompt ausführen", type="primary", width="stretch", disabled=run_disabled):
            if not api_key:
                st.toast("Kein API-Key für den gewählten Anbieter/Modell gespeichert (Seite: Einstellung).", icon="⚠️")
            else:
                if prov.startswith("OpenAI"):
                    results = run_custom_prompt_openai(custom_prompt.strip(), mdl, api_key, targets, base_df)
                else:
                    st.toast(f"Provider '{prov}' wird hier noch nicht unterstützt.", icon="⚠️")
                    results = {}

                if results:
                    st.session_state.base_df = _apply_custom_result_to_df(st.session_state.base_df, results)
                    st.toast(f"Ergebnisse für {len(results)} Bild(er) übernommen.", icon="✅")
                else:
                    st.toast("Kein verwertbares Ergebnis erhalten.", icon="⚠️")
        st.divider() 
    st.caption("Made with ♥️ by Elvis Owusu")

# ========= DB-Hilfen =========
def db_save_project(name: str, meta: dict, df: pd.DataFrame, images: Dict[str, bytes], thumbs: Dict[str, bytes], report_path: Optional[str] = None) -> None:
    data_json = json.dumps(df.to_dict(orient="records"), ensure_ascii=False)
    meta_json = json.dumps(meta, ensure_ascii=False)

    img_blob, th_blob = io.BytesIO(), io.BytesIO()
    with zipfile.ZipFile(img_blob, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for k, v in images.items(): z.writestr(k, v)
    with zipfile.ZipFile(th_blob, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for k, v in thumbs.items(): z.writestr(k, v)

    con = sqlite3.connect(DB_PATH); cur = con.cursor()
    cur.execute("""
        INSERT INTO projects (name,nr,addr,plz,ort,meta_json,data_json,images_zip,thumbs_zip,created_at,report_path)
        VALUES (?,?,?,?,?,?,?,?,?,?,?)
        ON CONFLICT(name) DO UPDATE SET
            nr=excluded.nr, addr=excluded.addr, plz=excluded.plz, ort=excluded.ort,
            meta_json=excluded.meta_json, data_json=excluded.data_json,
            images_zip=excluded.images_zip, thumbs_zip=excluded.thumbs_zip,
            created_at=excluded.created_at, report_path=excluded.report_path
    """, (name, meta.get("nr",""), meta.get("addr",""), meta.get("plz",""), meta.get("ort",""),
          meta_json, data_json, img_blob.getvalue(), th_blob.getvalue(),
          datetime.now().isoformat(), report_path or ""))
    con.commit(); con.close()

def db_list_projects_detailed() -> pd.DataFrame:
    con = sqlite3.connect(DB_PATH); cur = con.cursor()
    cur.execute("SELECT name, nr, addr, plz, ort, created_at, report_path FROM projects ORDER BY created_at DESC;")
    rows = cur.fetchall(); con.close()
    data=[]
    for name, nr, addr, plz, ort, created_at, report_path in rows:
        try: dt = datetime.fromisoformat(created_at).strftime("%d.%m.%Y %H:%M")
        except Exception: dt = created_at
        addr_full = ", ".join([part for part in [addr, f'{plz} {ort}'.strip()] if part and part.strip()])
        data.append({
            "Projektname": name or "", "Projektnummer": nr or "", "Adresse": addr_full or "",
            "Erstellungszeitpunkt": dt or "", "Bericht": report_path or ""
        })
    return pd.DataFrame(data)

def db_list_project_names() -> List[str]:
    con = sqlite3.connect(DB_PATH); cur = con.cursor()
    cur.execute("SELECT name FROM projects ORDER BY name COLLATE NOCASE;")
    names = [r[0] for r in cur.fetchall()]; con.close(); return names

def db_load_project(name: str) -> tuple[dict, pd.DataFrame, Dict[str, bytes], Dict[str, bytes]]:
    con = sqlite3.connect(DB_PATH); cur = con.cursor()
    cur.execute("SELECT meta_json,data_json,images_zip,thumbs_zip FROM projects WHERE name=?;", (name,))
    row = cur.fetchone(); con.close()
    if not row: raise ValueError("Projekt nicht gefunden.")
    meta = json.loads(row[0]); data = json.loads(row[1])
    images, thumbs = {}, {}
    with zipfile.ZipFile(io.BytesIO(row[2])) as z:
        for nm in z.namelist(): images[nm] = z.read(nm)
    with zipfile.ZipFile(io.BytesIO(row[3])) as z:
        for nm in z.namelist(): thumbs[nm] = z.read(nm)
    return meta, pd.DataFrame(data), images, thumbs

def db_delete_project(name: str):
    con = sqlite3.connect(DB_PATH); cur = con.cursor()
    cur.execute("DELETE FROM projects WHERE name=?;", (name,))
    con.commit(); con.close()

def reset_project_editor():
    st.session_state.meta = {"name":"", "nr":"", "addr":"", "plz":"", "ort":""}
    st.session_state.base_df = None
    st.session_state.images_map = {}
    st.session_state.thumbs_map = {}
    st.session_state.gallery_selected = None
    st.session_state.selected_for_ki = []
    st.session_state.last_preview_pdf = None

# ========= eBKP-Loader =========
@st.cache_data(show_spinner=False)
def load_ebkph_table(path: str) -> Optional[pd.DataFrame]:
    if not (path and os.path.exists(path)): return None
    try:
        with open(path, "rb") as f:
            data = f.read()
        buf = io.BytesIO(data)
        try:
            buf.seek(0)
            df = pd.read_csv(buf, sep=None, engine="python", encoding="utf-8", dtype=str).fillna("")
        except Exception:
            df = None
        if df is None or df.empty:
            for sep in [";", ",", "\t", "|"]:
                buf.seek(0)
                try:
                    df = pd.read_csv(buf, sep=sep, encoding="utf-8", dtype=str).fillna("")
                    if not df.empty: break
                except Exception:
                    continue
        if df is None or df.empty: return None
        cols_lower = {c.lower(): c for c in df.columns}
        code_col = cols_lower.get("code") or list(df.columns)[0]
        bez_col  = cols_lower.get("bezeichnung") or (list(df.columns)[1] if len(df.columns)>1 else list(df.columns)[0])
        df = df.rename(columns={code_col:"Code", bez_col:"Bezeichnung"})
        df["Code"] = df["Code"].astype(str).str.strip()
        df["Bezeichnung"] = df["Bezeichnung"].astype(str).str.strip()
        df = df[(df["Code"]!="") & (df["Bezeichnung"]!="")].drop_duplicates(subset=["Code","Bezeichnung"], keep="first").reset_index(drop=True)
        return df
    except Exception:
        return None

# ========= Filter =========
def apply_filters_and_sort(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty: return df
    f = st.session_state.filters
    view = df.copy()
    if f.get("geschoss"): view = view[view["Geschoss"].isin(f["geschoss"])]
    if f.get("kategorie"): view = view[view["Kategorie"].isin(f["kategorie"])]
    txt = (f.get("text") or "").strip()
    if txt:
        q = txt.lower()
        cols = [c for c in [
            "Bild original","Bild umbenennen","Bauteile","Besonderheiten","Risiken","Empfehlung",
            "Kategorie","Geschoss","Pfad"
        ] if c in view.columns]
        mask = pd.Series(False, index=view.index)
        for c in cols:
            mask = mask | view[c].astype(str).str.lower().str.contains(q, na=False)
        view = view[mask]
    return view


# ======== KI: Custom-Prompt auf ausgewählte Zeilen anwenden ========
def _get_active_ki_settings():
    # 1) Anbieter/Modell wie bisher aus der Config lesen
    cfg = st.session_state.config
    provider = cfg.get("ki_provider", "OpenAI (ChatGPT)")
    model = cfg.get("ki_model", "gpt-4o")

    # 2) API-Key bevorzugt aus Secrets/Umgebung (sicher!)
    secret_key = (st.secrets.get("OPENAI_API_KEY")
                  or os.environ.get("OPENAI_API_KEY")
                  or "")

    if secret_key:
        api_key = secret_key
    else:
        # 3) Fallback: alter Weg (nur falls noch irgendwo in config.json)
        api_keys = (cfg.get("api_keys") or {}).get(provider, {})
        api_key = api_keys.get(model) or api_keys.get("default") or ""

    return provider, model, api_key


def _apply_custom_result_to_df(df: pd.DataFrame, results: dict):
    """
    results = {
      "Bild original": {
        "Bauteile": "...",
        "Besonderheiten": "...",
        "Risiken": "...",
        "Empfehlung": "..."
      },
      ...
    }
    Schreibt die Felder in die DataFrame-Zeilen.
    """
    if df is None or df.empty: return df
    base = df.copy()
    for name, payload in results.items():
        mask = (base["Bild original"] == name) if "Bild original" in base.columns else None
        if mask is not None and mask.any():
            for col in ["Bauteile", "Besonderheiten", "Risiken", "Empfehlung"]:
                if col in base.columns and payload.get(col) is not None:
                    base.loc[mask, col] = str(payload.get(col, "")).strip()
    return base

def run_custom_prompt_openai(prompt: str, model: str, api_key: str, targets: list, df: pd.DataFrame) -> dict:
    """
    Ruft OpenAI (Chat Completions) auf und bittet um strukturiertes JSON mit
    Bauteile/Besonderheiten/Risiken/Empfehlung pro Bildname.
    """
    # Schutz: nichts tun, wenn keine Ziele
    if not targets: return {}

    # Kontext aus Tabelle (optional nützlich für das Modell)
    def _row_dict(name):
        try:
            row = df[df["Bild original"] == name].head(1).to_dict(orient="records")
            return row[0] if row else {"Bild original": name}
        except Exception:
            return {"Bild original": name}

    system_msg = (
        "Du bist ein Assistent für Zustandsanalysen in Gebäuden. "
        "Gib NUR JSON zurück. Struktur:\n"
        "{\n"
        '  "<bildname>": {\n'
        '    "Bauteile": "...",\n'
        '    "Besonderheiten": "...",\n'
        '    "Risiken": "...",\n'
        '    "Empfehlung": "..."\n'
        "  }, ...\n"
        "}"
    )

    user_payload = {
        "hinweis": "Überarbeite/ergänze die folgenden Datensätze.",
        "prompt": prompt,
        "ziele": [{ "Bild original": t, "context": _row_dict(t) } for t in targets]
    }

    # -------- OpenAI Client (chat.completions) --------
    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)}
            ],
            temperature=0.2,
        )
        content = resp.choices[0].message.content.strip()
    except Exception as e:
        st.toast(f"OpenAI-Aufruf fehlgeschlagen: {e}", icon="❌")
        return {}

    # JSON parsen (robust)
    try:
        # manchmal umschließt das Modell Codefences ```json ... ```
        if content.startswith("```"):
            content = content.strip("`")
            # nach optionalem "json" Kopfteil splitten
            if content.lower().startswith("json"):
                content = content[4:].lstrip()
        data = json.loads(content)
        if not isinstance(data, dict): raise ValueError("Antwort ist kein JSON-Objekt")
        results = {}
        for name in targets:
            block = data.get(name) or {}
            results[name] = {
                "Bauteile": block.get("Bauteile", ""),
                "Besonderheiten": block.get("Besonderheiten", ""),
                "Risiken": block.get("Risiken", ""),
                "Empfehlung": block.get("Empfehlung", ""),
            }
        return results
    except Exception as e:
        st.toast(f"Antwort nicht als JSON lesbar: {e}", icon="⚠️")
        return {}


# ========= PDF/EXCEL/ZIP Export =========
def _has_reportlab() -> bool:
    try:
        import reportlab  # noqa
        return True
    except Exception:
        return False

def export_pdf(df: pd.DataFrame, images_map: Dict[str, bytes], meta: dict) -> bytes:
    if not _has_reportlab():
        return b"%PDF-1.4\n% Minimaler Platzhalter-PDF. Installiere 'reportlab' fuer schoenes Layout.\n%%EOF"
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import cm
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4); w, h = A4

    # Deckblatt
    c.setFont("Helvetica-Bold", 18); c.drawString(2*cm, h-3*cm, f"Zustandsanalyse – {meta.get('name','')}")
    c.setFont("Helvetica", 11); y = h-4.2*cm
    for line in [
        f"Projektnummer: {meta.get('nr','')}",
        f"Adresse: {meta.get('addr','')} {meta.get('plz','')} {meta.get('ort','')}",
        f"Erstellt am: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
    ]:
        c.drawString(2*cm, y, line); y -= 0.6*cm
    c.showPage()

    cols = ["Bild original","Bild umbenennen","Zeit","Geschoss","Kategorie","Bauteile","Besonderheiten","Risiken","Empfehlung"]
    safe_df = df[[c_ for c_ in cols if c_ in df.columns]].copy()

    for _, r in safe_df.iterrows():
        c.setFont("Helvetica-Bold", 12)
        c.drawString(2*cm, h-2*cm, str(r.get("Bild umbenennen") or r.get("Bild original") or ""))
        c.setFont("Helvetica", 10); y = h-3*cm

        # Bild
        img_key = r.get("Bild original"); img_b = images_map.get(img_key)
        if img_b:
            try:
                from reportlab.lib.utils import ImageReader
                im = Image.open(io.BytesIO(img_b)); iw, ih = im.size
                img = ImageReader(io.BytesIO(img_b))
                maxw, maxh = 12*cm, 8*cm
                scale = min(maxw/iw, maxh/ih)
                c.drawImage(img, 2*cm, y-(ih*scale), width=iw*scale, height=ih*scale, preserveAspectRatio=True, mask='auto')
                y -= ih*scale + 0.7*cm
            except Exception:
                pass

        def draw(lbl, val):
            nonlocal y
            c.setFont("Helvetica-Bold", 10); c.drawString(2*cm, y, f"{lbl}:")
            c.setFont("Helvetica", 10); c.drawString(5*cm, y, str(val or "")); y -= 0.5*cm

        for lbl in ["Zeit","Geschoss","Kategorie","Bauteile","Besonderheiten","Risiken","Empfehlung"]:
            draw(lbl, r.get(lbl))
        c.showPage()
    c.save(); return buf.getvalue()

def export_excel(df: pd.DataFrame) -> bytes:
    out = io.BytesIO()
    try:
        with pd.ExcelWriter(out, engine="xlsxwriter") as writer:
            df.to_excel(writer, sheet_name="Zustandsanalyse", index=False)
            writer.close()
    except Exception:
        with pd.ExcelWriter(out) as writer:
            df.to_excel(writer, sheet_name="Zustandsanalyse", index=False)
            writer.close()
    return out.getvalue()

def export_zip(df: pd.DataFrame, images_map: Dict[str, bytes], pdf_bytes: bytes) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
        z.writestr("Bericht.pdf", pdf_bytes)
        if not df.empty and "Bild original" in df.columns:
            for img_name in df["Bild original"].dropna().unique():
                b = images_map.get(img_name)
                if b: z.writestr(f"bilder/{img_name}", b)
    return buf.getvalue()

# ========= WORD (DOCX) Export =========
def _has_docx() -> bool:
    try:
        import docx  # python-docx
        return True
    except Exception:
        return False

def export_word(df: pd.DataFrame, images_map: Dict[str, bytes], meta: dict) -> bytes:
    """
    Erzeugt einen Word-Bericht (DOCX) ähnlich wie das PDF:
    - Deckblatt mit Projektdaten
    - Pro Bild: Titel (umbenennen/Original), Bild (falls vorhanden) + Felder (Zeit, Geschoss, Kategorie, ...).
    """
    if not _has_docx():
        # Fallback: einfache Text-Datei als .docx verpackt (damit der Nutzer das Paket-Fehlen bemerkt)
        buf = io.BytesIO()
        buf.write(b"Bitte 'python-docx' installieren, um Word-Export zu aktivieren.\n")
        return buf.getvalue()

    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # Deckblatt
    title = doc.add_heading(level=0)
    run = title.add_run(f"Zustandsanalyse – {meta.get('name','')}")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    p.add_run(f"Projektnummer: {meta.get('nr','')}\n")
    p.add_run(f"Adresse: {meta.get('addr','')} {meta.get('plz','')} {meta.get('ort','')}\n")
    p.add_run(f"Erstellt am: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_page_break()

    # Spaltenauswahl (robust, wie im PDF)
    cols = ["Bild original","Bild umbenennen","Zeit","Geschoss","Kategorie","Bauteile","Besonderheiten","Risiken","Empfehlung"]
    safe_df = df[[c_ for c_ in cols if c_ in df.columns]].copy()

    for _, r in safe_df.iterrows():
        # Titel / Name
        hdr = doc.add_heading(level=1)
        hdr_run = hdr.add_run(str(r.get("Bild umbenennen") or r.get("Bild original") or ""))
        hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Bild (falls vorhanden)
        img_key = r.get("Bild original")
        img_b = images_map.get(img_key)
        if img_b:
            try:
                # Bild temporär schreiben (python-docx braucht Dateipfad oder Dateiobjekt)
                tmp = io.BytesIO(img_b)
                tmp.seek(0)
                # Breite auf ~12cm beschränken
                doc.add_picture(tmp, width=Cm(12))
            except Exception:
                pass

        # Felder
        def field(lbl: str, key: str):
            if key in r:
                p = doc.add_paragraph()
                p.add_run(f"{lbl}: ").bold = True
                p.add_run(str(r.get(key,"") or ""))

        field("Zeit", "Zeit")
        field("Geschoss", "Geschoss")
        field("Kategorie", "Kategorie")
        field("Bauteile", "Bauteile")
        field("Besonderheiten", "Besonderheiten")
        field("Risiken", "Risiken")
        field("Empfehlung", "Empfehlung")
        doc.add_page_break()

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

# ========= ZIP (ALLES) – Helpers =========
def _sanitize_folder_name(name: str) -> str:
    import re
    s = re.sub(r'[\\/:*?"<>|]+', "_", str(name or "Projekt")).strip()
    return s if s else "Projekt"

def _write_full_package_into_zip(z: zipfile.ZipFile, prefix: str, meta: dict,
                                 df: pd.DataFrame, images_map: Dict[str, bytes]) -> None:
    """
    Schreibt in ein bestehendes ZipFile-Objekt:
      <prefix>/Bericht.pdf
      <prefix>/Zustand.xlsx
      <prefix>/Bericht.docx (wenn python-docx installiert, sonst Hinweis)
      <prefix>/bilder/<alle Originalbilder>
    Erwartet, dass export_pdf/export_excel/export_word/_has_docx bereits vorhanden sind.
    """
    pdf_bytes = export_pdf(df, images_map, meta)
    excel_bytes = export_excel(df)
    docx_bytes = export_word(df, images_map, meta) if _has_docx() else None

    z.writestr(f"{prefix}/Bericht.pdf", pdf_bytes)
    z.writestr(f"{prefix}/Zustand.xlsx", excel_bytes)
    if docx_bytes:
        z.writestr(f"{prefix}/Bericht.docx", docx_bytes)
    else:
        z.writestr(f"{prefix}/HINWEIS_DOCUS.txt",
                   "python-docx ist nicht installiert – daher kein DOCX enthalten.")

    if not df.empty and ("Bild original" in df.columns):
        for img_name in df["Bild original"].dropna().unique():
            b = images_map.get(img_name)
            if b:
                z.writestr(f"{prefix}/bilder/{img_name}", b)


# ========= KI-Adapter (Demo, robust) =========
def run_ki_on_images(selected_names: List[str]) -> Dict[str, dict]:
    """
    Platzhalter-Analyse (strukturiert). Echte API-Calls können hier eingebaut werden.
    """
    results = {}
    for name in selected_names:
        results[name] = {
            "Bauteile": "Schaltkasten; Zuleitung; Steckdosen",
            "Besonderheiten": "Leichte Korrosion am Gehäuse; fehlende Beschriftung",
            "Risiken": "Überhitzung durch lockere Klemmen; Stolpergefahr durch lose Leitung",
            "Empfehlung": "Klemmen nachziehen; Beschriftung ergänzen; Leitung fixieren"
        }
    return results

# ========= Seiten =========

def page_projektspeicher(): #Seite Projekt
    """
    Einfache Projektliste + Suche + Löschen + Vorschau.
    Vorschau versteht:
    - lokalen PDF-Pfad
    - http(s)-Link (lädt PDF herunter)
    - data:application/pdf;base64,... (Base64 decodieren)
    Alles mit Edge-sicherer Ansicht + Download-Fallback.
    """
    st.subheader("Projektübersicht")

    # ---------- Suche ----------
    # Suche mit Icon links und Eingabefeld rechts
    col_icon, col_input = st.columns([0.08, 0.92])
    with col_icon:
        icon_path = os.path.join(BASE_DIR, "search.svg")
        if os.path.exists(icon_path):
            st.image(icon_path, width="stretch")
    with col_input:
        q = st.text_input(
            "Suche (Name / Nummer / Adresse)",
            value="", placeholder="Projektname, Projektnummer oder Adresse…",
            label_visibility="collapsed"  # Label ausgeblendet, da Icon schon da ist
        )
    
    st.divider()

    df = db_list_projects_detailed()
    if df.empty:
        st.info("Keine Projekte gespeichert.")
        return
    
    if q.strip():
        m = (
            df["Projektname"].str.lower().str.contains(q.lower(), na=False) |
            df["Projektnummer"].str.lower().str.contains(q.lower(), na=False) |
            df["Adresse"].str.lower().str.contains(q.lower(), na=False)
        )
        df = df[m]

    # ---------- Tabelle ----------
    df_view = df.copy()
    if "Auswählen" not in df_view.columns:
        df_view.insert(0, "Auswählen", False)

    edited = st.data_editor(
        df_view,
        width="stretch",
        height=500,
        column_config={
            "Auswählen": st.column_config.CheckboxColumn("Auswählen"),
            "Projektname": st.column_config.TextColumn(disabled=True),
            "Projektnummer": st.column_config.TextColumn(disabled=True),
            "Adresse": st.column_config.TextColumn(disabled=True),
            "Erstellungszeitpunkt": st.column_config.TextColumn(disabled=True),
            "Bericht": st.column_config.TextColumn(
                disabled=False,
                help="Pfad, http/https oder data:application/pdf;base64,..."
            ),
        },
        hide_index=True,
    )

    selected = edited.loc[edited["Auswählen"] == True, "Projektname"].tolist()

    # ---------- Löschen ----------
    c1, c2 = st.columns([0.5, 0.5])
    with c1:
        confirm = st.checkbox("Löschen bestätigen (nicht rückgängig)")
    with c2:
        if st.button("🗑️ Ausgewählte löschen", type="primary", disabled=not (selected and confirm), width="stretch"):
            try:
                for name in selected:
                    db_delete_project(name)
                st.success(f"{len(selected)} Projekt(e) gelöscht.")
                st.rerun()
            except Exception as e:
                st.error(f"Löschen fehlgeschlagen: {e}")

    # ---------- Vorschau ----------
    st.divider()
    st.subheader("Vorschau")

    if not selected:
        st.info("Bitte wähle ein Projekt für die Vorschau aus.")
        return

    first = selected[0]
    row = edited[edited["Projektname"] == first]
    report_path = row["Bericht"].values[0] if not row.empty else ""

    def _clean(s: str) -> str:
        s = (s or "").strip()
        if (s.startswith('"') and s.endswith('"')) or (s.startswith("'") and s.endswith("'")):
            s = s[1:-1].strip()
        return s

    pdf_bytes = None
    rp = _clean(str(report_path))

    try:
        if not rp:
            st.info("Kein Bericht angegeben.")
        elif rp.lower().startswith("data:application/pdf;base64,"):
            # data:-URL
            b64 = rp.split(",", 1)[1]
            pdf_bytes = base64.b64decode(b64)
        elif rp.lower().startswith(("http://", "https://")):
            # http(s) herunterladen
            try:
                import requests
                resp = requests.get(rp, timeout=10)
                resp.raise_for_status()
                pdf_bytes = resp.content
            except Exception as e:
                st.info(f"Download nicht möglich: {e}")
        elif os.path.exists(rp):
            # lokaler Pfad
            with open(rp, "rb") as f:
                pdf_bytes = f.read()
        else:
            st.info("Bericht nicht gefunden. Pfad/Link prüfen.")
    except Exception as e:
        st.info(f"Keine Inline-Vorschau möglich: {e}")

    if pdf_bytes:
        render_pdf_preview(pdf_bytes, height=840)
        st.download_button(
            "PDF herunterladen/öffnen",
            data=pdf_bytes,
            file_name=os.path.basename(rp) or f"{first}_Bericht.pdf",
            mime="application/pdf",
            width="stretch"
        )
    else:
        st.info("Kein gültiger Bericht für dieses Projekt vorhanden.")

def page_projekt(): #Seite Projekt erstellen
    """
    Einfache Erfassung + Tabelle + Vorschau.
    WICHTIG:
    - Vorschau funktioniert ohne KI: Button "Sofort-Vorschau" erzeugt PDF aus der aktuellen Tabelle.
    - Wenn später die KI läuft, kann sie ss.last_preview_pdf ebenfalls befüllen.
    """
    ss = st.session_state

    # ---------- Session-Defaults ----------
    for k, v in {
        "meta": {"name":"", "nr":"", "addr":"", "plz":"", "ort":""},
        "base_df": None,
        "images_map": {},
        "thumbs_map": {},
        "last_preview_pdf": None,
    }.items():
        if k not in ss: ss[k] = v

    st.subheader("Projekterfassung")

    # ----------Projekt laden-----------
    names = db_list_project_names()
    prev_selection = ss.get("load_selection", "")
    ss.load_selection = st.selectbox(
        "Gespeichertes Projekt laden",
        options=[""] + names,
        index=([""] + names).index(prev_selection) if prev_selection in ([""] + names) else 0,
    )
    selection = ss.load_selection
    if selection and selection != ss.loaded_project_name:
        try:
            meta, df, imgs, thumbs = db_load_project(selection)
            ss.meta = {"name": meta.get("name",""), "nr": meta.get("nr",""),
                       "addr": meta.get("addr",""), "plz": meta.get("plz",""), "ort": meta.get("ort","")}
            ss.base_df = df; ss.images_map = imgs; ss.thumbs_map = thumbs
            if "geschosse" in meta and isinstance(meta["geschosse"], list): ss.config["geschosse"] = meta["geschosse"]
            if "klassen" in meta and isinstance(meta["klassen"], list): ss.config["klassen"] = meta["klassen"]
            ss.loaded_project_name = selection
            st.toast(f"Projekt „{selection}“ geladen.", icon="✅")
        except Exception as e:
            st.toast(f"Projekt konnte nicht geladen werden: {e}", icon="❌")
    
    st.divider()

    # ---------- Speicher, Neu erstellen ----------
    c_new, c_discard, c_save = st.columns(3)
    with c_new:
        st.markdown("<div class='nowrap-btn'>", unsafe_allow_html=True)
        if st.button("Neu erfassen", type="secondary", width="stretch"):
            reset_project_editor(); ss.loaded_project_name = ""; ss.load_selection = ""
            st.toast("Neuerfassung gestartet. Alle Felder & Bilder wurden geleert.", icon="✨"); st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='nowrap-btn'>", unsafe_allow_html=True)
    run_ki = st.button("KI-Bauteilbeschrieb ausführen", width="stretch",
                       help="Analysiert markierte Bilder und füllt die Zustandsanalyse.")
    st.markdown("</div>", unsafe_allow_html=True)

    with c_discard:
        st.markdown("<div class='nowrap-btn'>", unsafe_allow_html=True)
        if st.button("Verwerfen", type="secondary", width="stretch"):
            reset_project_editor(); ss.loaded_project_name = ""; ss.load_selection = ""
            st.toast("Alle Inhalte verworfen (Eingaben & Bilder gelöscht).", icon="🗑️"); st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    with c_save:
        st.markdown("<div class='nowrap-btn'>", unsafe_allow_html=True)
        if st.button("Speichern", type="primary", width="stretch", help="Speichert Projekt, Daten & Bilder in der Datenbank."):
            pname = (ss.meta.get("name") or "").strip()
            pnr = (ss.meta.get("nr") or "").strip()
            if not pname: st.toast("„Projektname“ ist ein Pflichtfeld.", icon="⚠️")
            elif not pnr: st.toast("„Projektnummer“ ist ein Pflichtfeld.", icon="⚠️")
            elif ss.base_df is None or ss.base_df.empty or not ss.images_map:
                st.toast("Keine Daten/Bilder vorhanden.", icon="⚠️")
            else:
                try:
                    meta_to_save = {**ss.meta, "geschosse": ss.config["geschosse"], "klassen": ss.config["klassen"]}
                    db_save_project(pname, meta_to_save, ss.base_df, ss.images_map, ss.thumbs_map, report_path=None)
                    ss.load_selection = pname; ss.loaded_project_name = pname
                    st.toast(f"Projekt „{pname}“ erfolgreich gespeichert.", icon="✅")
                except Exception as e:
                    st.toast(f"Speichern fehlgeschlagen: {e}", icon="❌")
        st.markdown("</div>", unsafe_allow_html=True)

    if "pending_run_ki" not in ss: ss.pending_run_ki = False
    if run_ki: ss.pending_run_ki = True

    st.divider()

    # ---------- KI ausführen (wenn angefordert) ----------

    c1,c2,c3,c4,c5 = st.columns(5)
    with c1: ss.meta["name"] = st.text_input("Projektname", ss.meta.get("name",""))
    with c2: ss.meta["nr"]   = st.text_input("Projektnummer", ss.meta.get("nr",""))
    with c3: ss.meta["addr"] = st.text_input("Strasse / Nr.", ss.meta.get("addr",""))
    with c4: ss.meta["plz"]  = st.text_input("Postleitzahl", ss.meta.get("plz",""))
    with c5: ss.meta["ort"]  = st.text_input("Ort", ss.meta.get("ort",""))

    st.divider()

    # ---------- Upload ----------
    st.write("**Bilder importieren (JPG/PNG/ZIP)**")
    files = st.file_uploader(
        "Dateien auswählen",
        type=["jpg","jpeg","png","zip"],
        accept_multiple_files=True
    )

    # Mini-Helper für neuen Tabellen-Eintrag
    def _add_record(fname: str, data: bytes, pfad: str):
        from PIL import Image
        img = Image.open(io.BytesIO(data)).convert("RGB")
        th = img.copy(); th.thumbnail((256, 256))
        bth = io.BytesIO(); th.save(bth, format="JPEG", quality=80)
        rec = {
            "Miniatur": f"data:image/jpeg;base64,{base64.b64encode(bth.getvalue()).decode('ascii')}",
            "Bild original": fname,
            "Bild umbenennen": "",
            "Zeit": "",
            "Geschoss": "Unbekannt",
            "Kategorie": "Sonstiges / Umfeld",
            "Bauteile": "",
            "Besonderheiten": "",
            "Risiken": "",
            "Empfehlung": "",
            "Pfad": pfad,
        }
        return rec, bth.getvalue()

    if files:
        new_rows = []
        new_imgs = {}
        new_thumbs = {}
        added = 0

        for f in files:
            if f.name.lower().endswith(".zip"):
                with zipfile.ZipFile(io.BytesIO(f.read())) as z:
                    for nm in z.namelist():
                        if nm.lower().endswith((".jpg",".jpeg",".png")):
                            data = z.read(nm)
                            if nm not in ss.images_map:
                                rec, th = _add_record(os.path.basename(nm), data, pfad=nm)
                                new_rows.append(rec)
                                new_imgs[os.path.basename(nm)] = data
                                new_thumbs[os.path.basename(nm)] = th
                                added += 1
            else:
                data = f.read()
                if f.name not in ss.images_map:
                    rec, th = _add_record(f.name, data, pfad=f.name)
                    new_rows.append(rec)
                    new_imgs[f.name] = data
                    new_thumbs[f.name] = th
                    added += 1

        if new_rows:
            df_new = pd.DataFrame(new_rows)
            if ss.base_df is None or ss.base_df.empty:
                ss.base_df = df_new
            else:
                ss.base_df = pd.concat([ss.base_df, df_new], ignore_index=True)
                ss.base_df = ss.base_df.drop_duplicates(subset=["Bild original"], keep="first").reset_index(drop=True)
            ss.images_map.update(new_imgs)
            ss.thumbs_map.update(new_thumbs)

        if added:
            st.success(f"{added} Bild(er) hinzugefügt.")

    # ---------- Tabelle ----------
    if ss.base_df is not None and not ss.base_df.empty:
        st.subheader("Zustandsanalyse (bearbeitbar)")
        edited = st.data_editor(
            ss.base_df,
            width="stretch",
            height=520,
            column_config={
                "Miniatur": st.column_config.ImageColumn("Miniatur", width="small"),
                "Bild original": st.column_config.TextColumn(disabled=True),
                "Bild umbenennen": st.column_config.TextColumn(),
                "Zeit": st.column_config.TextColumn(),
                "Geschoss": st.column_config.SelectboxColumn(options=ss.config["geschosse"]),
                "Kategorie": st.column_config.SelectboxColumn(options=ss.config["klassen"]),
                "Bauteile": st.column_config.TextColumn(width="large"),
                "Besonderheiten": st.column_config.TextColumn(width="large"),
                "Risiken": st.column_config.TextColumn(width="large"),
                "Empfehlung": st.column_config.TextColumn(width="large"),
                "Pfad": st.column_config.TextColumn(disabled=True, width="large"),
            },
        )
        ss.base_df = edited
    else:
        st.info("Noch keine Daten vorhanden. Lade Bilder hoch.")

    st.divider()
    st.subheader("Galerie")

#----------------- Galerie-----------------
    view_now = apply_filters_and_sort(ss.base_df)
    cards_per_row = 4
    for i in range(0, len(view_now), cards_per_row):
        chunk = view_now.iloc[i:i+cards_per_row]
        cols = st.columns(cards_per_row)
        for col, (_, r) in zip(cols, chunk.iterrows()):
            with col:
                orig = r["Bild original"]
                title = r["Bild umbenennen"].strip() if str(r["Bild umbenennen"]).strip() else orig
                tbytes = ss.thumbs_map.get(orig)
                if tbytes: st.image(tbytes, width="stretch")
                st.caption(f"**{title}**")
                if st.button("Details", key=f"det_{orig}"): ss.gallery_selected = orig

    sel = ss.gallery_selected
    if sel:
        st.divider(); st.markdown("### Details")
        img_b = ss.images_map.get(sel)
        if img_b: st.image(img_b, caption=sel, width="stretch")
        row = ss.base_df[ss.base_df["Bild original"] == sel]
        if not row.empty:
            r = row.iloc[0]
            c1,c2,c3 = st.columns(3)
            with c1: st.write("**Zeit:**", r.get("Zeit",""))
            with c2: st.write("**Geschoss:**", r.get("Geschoss",""))
            with c3: st.write("**Kategorie:**", r.get("Kategorie",""))
            st.write("**Bauteile:**", r.get("Bauteile",""))
            st.write("**Besonderheiten:**", r.get("Besonderheiten",""))
            st.write("**Risiken:**", r.get("Risiken",""))
            st.write("**Empfehlung:**", r.get("Empfehlung",""))

    st.divider()
    st.subheader("Vorschau")

    # ---------- Sofort-Vorschau ----------
    col_btn, _ = st.columns([0.35, 0.65])
    with col_btn:
        if st.button("Sofort-Vorschau aus aktueller Tabelle erzeugen", width="stretch"):
            try:
                if ss.base_df is None or ss.base_df.empty:
                    st.warning("Keine Daten für die Vorschau.")
                else:
                    pdf_bytes = export_pdf(ss.base_df.copy(), ss.images_map, ss.meta)
                    ss.last_preview_pdf = pdf_bytes
                    st.success("Vorschau erzeugt.")
            except Exception as e:
                st.error(f"Vorschau konnte nicht erzeugt werden: {e}")

    # ---------- Anzeige + Download ----------
    if ss.last_preview_pdf:
        render_pdf_preview(ss.last_preview_pdf, height=840)
        st.download_button(
            "PDF herunterladen/öffnen",
            data=ss.last_preview_pdf,
            file_name=(ss.meta.get("name") or "Projekt") + "_Vorschau.pdf",
            mime="application/pdf",
            width="stretch"
        )
    else:
        st.info("Noch keine Vorschau vorhanden (Sofort-Vorschau drücken).")

def page_export():
    """
    Exportbereich:
    - Suche + Mehrfachauswahl von Projekten
    - Einzel-/Mehrfach-Export:
        * PDF (Einzel direkt, Mehrfach als ZIP)
        * Word (DOCX) (Einzel direkt, Mehrfach als ZIP)
        * Excel (Einzel direkt, Mehrfach als ZIP)
        * ZIP (pro Projekt ein Ordner mit PDF+DOCX+Excel+Bilder)
    """
    st.subheader("Exportbereich")

    # Suche (mit optionalem Icon)
    sc, tc = st.columns([0.08, 0.92])
    with sc:
        icon_path = os.path.join(BASE_DIR, "search.svg")
        if os.path.exists(icon_path):
            st.image(icon_path, width="stretch")
    with tc:
        query = st.text_input("Suche", value="", placeholder="Projektname, Projektnummer oder Adresse…",
                              label_visibility="collapsed")

    # Projekte laden + filtern
    df_list = db_list_projects_detailed()
    if df_list.empty:
        st.info("Keine Projekte gespeichert.")
        return

    if query.strip():
        q = query.lower()
        mask = (
            df_list["Projektname"].str.lower().str.contains(q, na=False) |
            df_list["Projektnummer"].str.lower().str.contains(q, na=False) |
            df_list["Adresse"].str.lower().str.contains(q, na=False)
        )
        df_list = df_list[mask]

    st.divider()

    # Selektions-Tabelle
    df_view = df_list.copy()
    if "Auswählen" not in df_view.columns:
        df_view.insert(0, "Auswählen", False)

    edited = st.data_editor(
        df_view,
        width="stretch",
        height=500,
        column_config={
            "Auswählen": st.column_config.CheckboxColumn("Auswählen"),
            "Projektname": st.column_config.TextColumn(disabled=True),
            "Projektnummer": st.column_config.TextColumn(disabled=True),
            "Adresse": st.column_config.TextColumn(disabled=True),
            "Erstellungszeitpunkt": st.column_config.TextColumn(disabled=True),
            "Bericht": st.column_config.TextColumn(disabled=False, help="Pfad oder Link zu einer PDF-Datei für die Vorschau"),
        },
        hide_index=True,
    )
    selected = edited.loc[edited["Auswählen"] == True, "Projektname"].tolist()

    st.divider()
    export_type = st.radio(
        "Exporttyp wählen",
        options=["PDF Bericht", "Word (DOCX)", "Excel", "ZIP (PDF+DOCX+Excel+Bilder)"],
        horizontal=True
    )
    st.divider()

    if not selected:
        st.info("Bitte wähle mindestens ein Projekt aus.")
        return

    # Helper: sicherer Ordnername (verwende vorhandene Helper-Funktion, falls im File vorhanden)
    def safe_name(name: str) -> str:
        try:
            return _sanitize_folder_name(name)
        except Exception:
            return re.sub(r'[\\/:*?"<>|]+', "_", str(name or "Projekt")).strip().strip(".")

    # ========== PDF ==========
    if export_type == "PDF Bericht":
        if len(selected) == 1:
            first = selected[0]
            try:
                meta, df_data, images_map, thumbs_map = db_load_project(first)
            except Exception as e:
                st.error(f"Projekt konnte nicht geladen werden: {e}")
                return

            pdf_bytes = export_pdf(df_data.copy(), images_map, meta)
            st.download_button("📄 PDF herunterladen",
                               data=pdf_bytes,
                               file_name=f"{safe_name(first)}_Bericht.pdf",
                               mime="application/pdf")

            st.subheader("Vorschau")
            render_pdf_preview(pdf_bytes, height=840)

            # Optional: aktuellen Berichtspfad ins Projekt schreiben
            if st.checkbox("Als aktuellen Berichtspfad im Projekt speichern"):
                out_dir = os.path.join(BASE_DIR, "berichte")
                os.makedirs(out_dir, exist_ok=True)
                fpath = os.path.join(out_dir, f"{safe_name(first)}_Bericht.pdf")
                try:
                    with open(fpath, "wb") as f:
                        f.write(pdf_bytes)
                    meta_to_save = {
                        **meta,
                        "geschosse": st.session_state.config["geschosse"],
                        "klassen": st.session_state.config["klassen"],
                    }
                    db_save_project(first, meta_to_save, df_data, images_map, thumbs_map, report_path=fpath)
                    st.success("Berichtspfad im Projekt gespeichert.")
                except Exception as e:
                    st.error(f"Speichern fehlgeschlagen: {e}")
        else:
            # Mehrere Projekte -> ZIP mit einzelnen PDFs
            buf = io.BytesIO()
            with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
                for pname in selected:
                    try:
                        meta, df_data, images_map, _ = db_load_project(pname)
                        pdf_bytes = export_pdf(df_data.copy(), images_map, meta)
                        z.writestr(f"{safe_name(pname)}_Bericht.pdf", pdf_bytes)
                    except Exception as e:
                        z.writestr(f"{safe_name(pname)}_FEHLER.txt", f"PDF-Export fehlgeschlagen: {e}")
            st.download_button("🧩 PDFs als ZIP herunterladen",
                               data=buf.getvalue(),
                               file_name=f"{safe_name(selected[0])}_und_weitere_PDF.zip",
                               mime="application/zip")

    # ========== DOCX ==========
    elif export_type == "Word (DOCX)":
        if not _has_docx():
            st.error("python-docx ist nicht installiert. Bitte mit 'pip install python-docx' nachrüsten.")
            return

        if len(selected) == 1:
            first = selected[0]
            try:
                meta, df_data, images_map, _ = db_load_project(first)
            except Exception as e:
                st.error(f"Projekt konnte nicht geladen werden: {e}")
                return

            docx_bytes = export_word(df_data.copy(), images_map, meta)
            st.download_button("📝 Word (DOCX) herunterladen",
                               data=docx_bytes,
                               file_name=f"{safe_name(first)}_Bericht.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            buf = io.BytesIO()
            with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
                for pname in selected:
                    try:
                        meta, df_data, images_map, _ = db_load_project(pname)
                        docx_bytes = export_word(df_data.copy(), images_map, meta)
                        z.writestr(f"{safe_name(pname)}_Bericht.docx", docx_bytes)
                    except Exception as e:
                        z.writestr(f"{safe_name(pname)}_FEHLER.txt", f"DOCX-Export fehlgeschlagen: {e}")
            st.download_button("🧩 DOCX als ZIP herunterladen",
                               data=buf.getvalue(),
                               file_name=f"{safe_name(selected[0])}_und_weitere_DOCX.zip",
                               mime="application/zip")

    # ========== Excel ==========
    elif export_type == "Excel":
        if len(selected) == 1:
            first = selected[0]
            try:
                _, df_data, _, _ = db_load_project(first)
            except Exception as e:
                st.error(f"Projekt konnte nicht geladen werden: {e}")
                return

            excel_bytes = export_excel(df_data.copy())
            st.download_button("📊 Excel herunterladen",
                               data=excel_bytes,
                               file_name=f"{safe_name(first)}_Zustand.xlsx",
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            buf = io.BytesIO()
            with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
                for pname in selected:
                    try:
                        _, df_data, _, _ = db_load_project(pname)
                        excel_bytes = export_excel(df_data.copy())
                        z.writestr(f"{safe_name(pname)}_Zustand.xlsx", excel_bytes)
                    except Exception as e:
                        z.writestr(f"{safe_name(pname)}_FEHLER.txt", f"Excel-Export fehlgeschlagen: {e}")
            st.download_button("🧩 Excel-Dateien als ZIP herunterladen",
                               data=buf.getvalue(),
                               file_name=f"{safe_name(selected[0])}_und_weitere_Excel.zip",
                               mime="application/zip")

    # ========== ZIP: PDF + DOCX + Excel + Bilder ==========
    elif export_type == "ZIP (PDF+DOCX+Excel+Bilder)":
        # Immer pro Projekt ein Unterordner -> vermeidet Probleme mit Pfaden wie "." oder führenden Slashes
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
            for pname in selected:
                try:
                    meta, df_data, images_map, _ = db_load_project(pname)
                    prefix = safe_name(pname)  # NIE leer lassen -> sonst fehlerhafte Zip-Pfade
                    _write_full_package_into_zip(z, prefix, meta, df_data.copy(), images_map)
                except Exception as e:
                    z.writestr(f"{safe_name(pname)}/FEHLER.txt", f"Export fehlgeschlagen: {e}")

        st.download_button("🗜️ Komplett-ZIP herunterladen (PDF+DOCX+Excel+Bilder)",
                           data=buf.getvalue(),
                           file_name=f"{safe_name(selected[0])}_und_weitere_Pakete.zip",
                           mime="application/zip")

def page_einstellung():
    # NEU: Geschosse verwalten (vor Kategorien)
    st.markdown("### Geschosse verwalten")
    current_geschosse = st.session_state.config.get("geschosse", DEFAULT_GESCHOSSE)
    geschosse_text = st.text_area(
        "Geschosse (eine Zeile pro Eintrag, Reihenfolge = Auswahlsortierung)",
        value="\n".join(current_geschosse),
        height=140,
        help="Beispiele: UG, EG, 1. OG, 2. OG, DG, Attika …"
    )
    if st.button("Geschosse speichern", type="primary"):
        new_geschosse = [x.strip() for x in geschosse_text.splitlines() if x.strip()]
        if new_geschosse:
            st.session_state.config["geschosse"] = new_geschosse
            save_config(st.session_state.config)
            st.success("Geschosse gespeichert. Sie stehen ab sofort in Filtern und in der Tabelle zur Auswahl.")

    st.markdown("---")
    st.markdown("### Kategorien verwalten")
    current_klassen = st.session_state.config.get("klassen", DEFAULT_KLASSEN)
    klassen_text = st.text_area(
        "Kategorien (eine pro Zeile)",
        value="\n".join(current_klassen),
        height=160,
        help="Eigene Kategorien definieren. Diese stehen in der Tabelle zur Auswahl und werden gespeichert."
    )
    if st.button("Kategorien speichern", type="primary"):
        new_klassen = [x.strip() for x in klassen_text.splitlines() if x.strip()]
        if new_klassen:
            st.session_state.config["klassen"] = new_klassen
            save_config(st.session_state.config)
            st.success("Kategorien gespeichert.")

    # ========= eBKP-H Auswahl mit Tree + Speichern/Laden =========
    st.markdown("---")
    st.markdown("### eBKP-H Auswahl")

    df_ebk = load_ebkph_table(EBKPH_CSV_PATH)
    if df_ebk is None or df_ebk.empty:
        st.info("CSV/Excel nicht gefunden oder leer. Lege eine gültige eBKP-H Datei am Standardpfad ab."); return

    # Baum aufbauen
    def _normalize_code(code: str) -> str:
        if not isinstance(code, str): code = str(code or "")
        code = code.strip()
        code = code.replace("•",".").replace("·",".").replace("‧",".").replace("∙",".").replace("・",".")
        code = re.sub(r"\s+", "", code); return code

    df_ebk = df_ebk.copy()
    df_ebk["Code"] = df_ebk["Code"].map(_normalize_code)
    df_ebk = df_ebk[(df_ebk["Code"]!="") & (df_ebk["Bezeichnung"]!="")]
    code_to_bez = dict(zip(df_ebk["Code"], df_ebk["Bezeichnung"]))

    def insert_into_tree(tree: dict, raw_code: str, bezeichnung: str):
        code = _normalize_code(raw_code)
        if not code: return
        parts = [p for p in code.split('.') if p]; prefixes: List[str] = []
        if not parts: return
        first = parts[0]; m = re.match(r"^[A-Za-z]+", first)
        if m:
            prefixes.append(m.group(0))
            if len(first) > len(m.group(0)): prefixes.append(first)
        else:
            prefixes.append(first)
        cur = prefixes[-1]
        for part in parts[1:]:
            cur = f"{cur}.{part}"; prefixes.append(cur)
        node = tree
        for i, key in enumerate(prefixes):
            if "children" not in node: node["children"] = {}
            if key not in node["children"]:
                node["children"][key] = {"code": key, "bez": bezeichnung if i==len(prefixes)-1 else key, "children": {}}
            node = node["children"][key]
            if i == len(prefixes)-1: node["bez"] = bezeichnung

    def build_tree(df: pd.DataFrame) -> dict:
        t = {"code":"root","bez":"root","children":{}}
        for _, r in df.iterrows(): insert_into_tree(t, r["Code"], r["Bezeichnung"])
        return t

    tree = build_tree(df_ebk)
    if "tree_init_done" not in st.session_state: st.session_state.tree_init_done = False
    if not st.session_state.tree_init_done:
        st.session_state.tree_init_done = True

    # Sortierfunktion für Codes
    def sort_key(code: str):
        parts = code.split(".")
        key = []
        for p in parts:
            if p.isdigit():
                key.append((0, int(p)))
            else:
                key.append((1, p))
        return key

    # Utilities
    def get_all_codes(node: dict, dest: List[str]):
        if node["code"] != "root":
            dest.append(node["code"])
        for ch in node.get("children", {}).values():
            get_all_codes(ch, dest)

    def get_selected_codes(node: dict, dest: List[str]):
        if node["code"] != "root":
            if st.session_state.get(f"tree_{node['code']}", False):
                dest.append(node["code"])
        for ch in node.get("children", {}).values():
            get_selected_codes(ch, dest)

    # --- Pending Bulk-Aktionen vor Rendern anwenden
    all_codes: List[str] = []
    get_all_codes(tree, all_codes)
    pending = st.session_state.pop("ebkh_apply", None)
    if pending:
        op = pending.get("op")
        if op == "select_all":
            for code in all_codes:
                st.session_state[f"tree_{code}"] = True
        elif op == "clear_all":
            for code in all_codes:
                st.session_state[f"tree_{code}"] = False
        elif op == "load_last":
            rows = set(pending.get("codes", []))
            for code in all_codes:
                st.session_state[f"tree_{code}"] = (code in rows)

    # Renderer
    def render_node(node: Dict[str,Any], level: int = 0):
        if node["code"] == "root":
            for child_key in sorted(node["children"].keys(), key=sort_key):
                render_node(node["children"][child_key], 0)
            return
        code = node["code"]; bez = code_to_bez.get(code, node.get("bez",""))
        label = f"{code} – {bez}"; children = node.get("children", {})
        if level == 0:
            with st.expander(label, expanded=False):
                for child_key in sorted(children.keys(), key=sort_key):
                    render_node(children[child_key], 1)
            return
        if level == 1:
            with st.expander(label, expanded=False):
                for child_key in sorted(children.keys(), key=sort_key):
                    render_node(children[child_key], 2)
            # Level-1: Checkbox + Label (keine Abschnitt-Buttons mehr)
            cols = st.columns([0.06, 0.94])
            #with cols[0]: st.checkbox("", key=f"tree_{code}")
            #with cols[1]: st.markdown(f"**{label}**")
            return
        # Level-2 (und tiefer): NUR Label, KEINE Checkbox
        # Level-2: Checkbox + Label
        if level == 2:
            cols = st.columns([0.06, 0.94])
            with cols[0]:
                st.checkbox("", key=f"tree_{code}")
            with cols[1]:
                st.markdown(label)
            return
        # Level >= 3: weiterhin nur Label
        st.markdown(label)

    # Jetzt erst UI rendern
    render_node(tree)

    # Zusammenfassung + Buttons
    selected_codes: List[str] = []
    get_selected_codes(tree, selected_codes)
    st.caption(f"Aktuell markiert: **{len(selected_codes)}** Codes")

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        if st.button("Alles auswählen"):
            st.session_state["ebkh_apply"] = {"op": "select_all"}
            st.rerun()
    with c2:
        if st.button("Alles leeren"):
            st.session_state["ebkh_apply"] = {"op": "clear_all"}
            st.rerun()
    with c3:
        if st.button("Auswahl speichern", type="primary"):
            try:
                if not selected_codes:
                    st.toast("Bitte zuerst Codes auswählen.", icon="⚠️")
                else:
                    ts = datetime.now().isoformat(timespec="seconds")
                    con = sqlite3.connect(DB_PATH); cur = con.cursor()
                    cur.executemany("INSERT INTO ebkh_selection (code, batch_ts) VALUES (?,?)",
                                    [(c, ts) for c in selected_codes])
                    con.commit(); con.close()
                    st.toast(f"Auswahl gespeichert ({len(selected_codes)} Codes, {ts}).", icon="✅")
            except Exception as e:
                st.toast(f"Speichern fehlgeschlagen: {e}", icon="❌")
    with c4:
        if st.button("Letzte Auswahl laden"):
            con = sqlite3.connect(DB_PATH); cur = con.cursor()
            cur.execute("SELECT MAX(batch_ts) FROM ebkh_selection")
            ts = cur.fetchone()[0]
            if not ts:
                con.close()
                st.toast("Keine gespeicherte Auswahl gefunden.", icon="⚠️")
            else:
                cur.execute("SELECT code FROM ebkh_selection WHERE batch_ts=?", (ts,))
                rows = [r[0] for r in cur.fetchall()]
                con.close()
                if not rows:
                    st.toast("Keine gespeicherten Codes im letzten Batch.", icon="⚠️")
                else:
                    st.session_state["ebkh_apply"] = {"op": "load_last", "codes": rows}
                    st.rerun()

    st.markdown("---")
    st.markdown("### KI-Anbieter & Modell")

    # Modelle je Anbieter
    PROVIDER_MODELS = {
        "OpenAI (ChatGPT)": ["gpt-5","gpt-4.1","gpt-4o","gpt-4-turbo","gpt-3.5-turbo"],
        "Microsoft Copilot": ["copilot-gpt-5","copilot-gpt-4.1","copilot-gpt-4o","copilot-gpt-4-turbo","copilot-gpt-35-turbo"],
    }

    # Aktuelle Auswahl laden
    prov = st.selectbox(
        "KI-Anbieter",
        list(PROVIDER_MODELS.keys()),
        index=0 if st.session_state.config.get("ki_provider","").startswith("OpenAI") else 1,
        key="ki_provider"
    )

    models = PROVIDER_MODELS.get(prov, [])
    mdl = st.selectbox(
        "Modell-Version",
        models,
        index=(models.index(st.session_state.config.get("ki_model")) if st.session_state.config.get("ki_model") in models else 0),
        key="ki_model"
    )

    # Sicherstellen, dass der Key-Store existiert
    if "api_keys" not in st.session_state.config or not isinstance(st.session_state.config["api_keys"], dict):
        st.session_state.config["api_keys"] = {"OpenAI (ChatGPT)": {}, "Microsoft Copilot": {}}
    if prov not in st.session_state.config["api_keys"]:
        st.session_state.config["api_keys"][prov] = {}

    # Gespeicherten Key für diese Kombi ermitteln (Fallback: "default")
    stored_key = (
        st.session_state.config["api_keys"][prov].get(mdl)
        or st.session_state.config["api_keys"][prov].get("default", "")
    )

    # WICHTIG: dynamischer Widget-Key, damit beim Wechsel von Anbieter/Modell das Feld neu gebildet wird
    pw_widget_key = f"api_key_widget::{prov}::{mdl}"

    api_key_input = st.text_input(
        f"{prov} API-Key (für {mdl})",
        type="password",
        value=stored_key,     # vorbefüllt (verdeckt) falls vorhanden
        key=pw_widget_key
    )

    c1, c2 = st.columns(2)
    with c1:
        if st.button("KI-Einstellungen speichern", type="primary"):
            # Nur speichern, wenn etwas eingegeben wurde (leerer String überschreibt NICHT)
            entered = st.session_state.get(pw_widget_key, "")
            if entered.strip():
                st.session_state.config["api_keys"][prov][mdl] = entered.strip()
            # Provider/Modell-Auswahl selbst merken
            save_config(st.session_state.config)
            st.success(f"Einstellungen gespeichert für: {prov} – {mdl}")

    with c2:
        if st.button("Key löschen"):
            # Entfernt nur den Key für genau diese Kombi
            try:
                if mdl in st.session_state.config["api_keys"][prov]:
                    st.session_state.config["api_keys"][prov].pop(mdl, None)
                    save_config(st.session_state.config)
                    # UI-Feld leeren, indem wir den Widget-State zurücksetzen
                    st.session_state[pw_widget_key] = ""
                    st.success(f"API-Key für {prov} / {mdl} gelöscht.")
                else:
                    st.info("Für diese Kombination ist kein Key gespeichert.")
            except Exception as e:
                st.error(f"Konnte Key nicht löschen: {e}")

    st.caption("Hinweis: Das Feld zeigt gespeicherte Schlüssel immer verdeckt an. "
            "Zum Aktualisieren einfach einen neuen Key eingeben und speichern.")


# ========== Zeigt ein PDF sicher im Browser an ==========

# ========== PDF Inline-Preview Helper ==========
def render_pdf_preview(pdf_bytes: bytes | None, height: int = 840):
    """Zeigt ein PDF (als Bytes) inline an – Edge-sicher via Blob-URL."""
    import streamlit as st
    import base64
    if not pdf_bytes:
        st.info("Noch keine Vorschau vorhanden.")
        return

    b64 = base64.b64encode(pdf_bytes).decode("ascii")
    html = f"""
    <div id="pdf-root" style="width:100%;height:{height}px;border:1px solid #e5e5e5;border-radius:8px;overflow:hidden">
      <iframe id="pdf-frame" style="width:100%;height:100%;border:0" sandbox="allow-scripts allow-same-origin"></iframe>
    </div>
    <script>
      (function(){{
        try {{
          const binStr = atob("{b64}");
          const len = binStr.length;
          const bytes = new Uint8Array(len);
          for (let i=0;i<len;i++) bytes[i] = binStr.charCodeAt(i);
          const blob = new Blob([bytes], {{type:"application/pdf"}});
          const url  = URL.createObjectURL(blob);
          const f = document.getElementById("pdf-frame");
          f.src = url + "#toolbar=1";
        }} catch(err) {{
          document.getElementById("pdf-root").innerHTML =
            "<div style='padding:12px;color:#b00020'>Inline-Vorschau nicht möglich. Bitte Datei herunterladen.</div>";
          console.error(err);
        }}
      }})();
    </script>
    """
    st.components.v1.html(html, height=height+20)




# ========= Router =========
page = st.session_state.page
if page == "Projekt":
    page_projekt()
elif page == "Projektspeicher":
    page_projektspeicher()
elif page == "Export":
    page_export()
else:
    page_einstellung()
